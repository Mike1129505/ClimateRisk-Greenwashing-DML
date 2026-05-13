import pandas as pd
import numpy as np
import os
import warnings
from sklearn.ensemble import RandomForestRegressor
from sklearn.linear_model import LassoCV
from sklearn.model_selection import GridSearchCV
from sklearn.preprocessing import StandardScaler
from doubleml import DoubleMLData, DoubleMLPLR
from econml.dml import CausalForestDML
from scipy.stats import pearsonr, norm
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings('ignore')

# ---------- 1. Environment & Parameter Settings ----------
#data_folder = #please enter folder
#file_path = os.path.join(data_folder, '.xlsx')
#df_raw = pd.read_excel(file_path).dropna().reset_index(drop=True)

T_vars = ['Trans_Risk']
Y_vars = ['GW_Score']
M_vars = ['Managerial_Myopia', 'Opacity']  # Mediator pool
Moderator = 'CEO_Risk_Pref'
X_vars = [
    'Size', 'Lev', 'ROA', 'Digitalization', 'Cashflow', 'Liquid', 'Quick', 'Tang_Ratio',
    'Intang_Ratio', 'Growth', 'Asset_Growth', 'Fin_Constraint', 'TobinQ', 'BM',
    'Analyst', 'Board_Size', 'Indep_Ratio', 'Inst_Hold', 'Mng_Hold', 'Fin_Back',
    'Oversea_Back', 'Big4', 'Wind_E_Score', 'TFP', 'EPS', 'RD_Invest',
    'Subsidy', 'HHI', 'IC_Quality', 'Media_Neg', 'Wage', 'List_Age', 'Firm_Age'
]

# ---------- 2. Fixed Effects (FE): Within-Transformation ----------
print(">>> Processing Fixed Effects (Entity & Time demeaning)...")
df = df_raw.copy()
all_cols = X_vars + T_vars + Y_vars + M_vars + [Moderator]
for col in all_cols:
    if np.issubdtype(df[col].dtype, np.number):
        # Subtract Entity mean and Year mean to handle FE
        df[col] = df[col] - df.groupby('Stkcd')[col].transform('mean')
        df[col] = df[col] - df.groupby('Year')[col].transform('mean')

# ---------- 3. Double Selection (Lasso Variable Selection) ----------
print(">>> Performing Double Selection to identify key controls...")


def double_selection(data, y_name, t_name, x_pool):
    lasso_y = LassoCV(cv=5).fit(data[x_pool], data[y_name])
    lasso_t = LassoCV(cv=5).fit(data[x_pool], data[t_name])
    selected = set(data[x_pool].columns[np.abs(lasso_y.coef_) > 1e-6]) | \
               set(data[x_pool].columns[np.abs(lasso_t.coef_) > 1e-6])
    return list(selected)


selected_X = double_selection(df, Y_vars[0], T_vars[0], X_vars)
# ---------- 4. Hyperparameter Tuning (Grid Search) ----------
print(">>> Tuning Random Forest hyperparameters for DML...")
rf_params = {
    'n_estimators': [300, 500],
    'max_depth': [5, 8, 10],
    'min_samples_leaf': [5, 10]
}
grid_search = GridSearchCV(RandomForestRegressor(random_state=42), rf_params, cv=5, scoring='r2', n_jobs=-1)
grid_search.fit(df[selected_X], df[Y_vars[0]])
best_params = grid_search.best_params_
print(f"Best Hyperparameters: {best_params}")

# ---------- 5. Orthogonality Diagnostics ----------
print("\n>>> Running Orthogonality Check...")
dml_data = DoubleMLData(df, y_col=Y_vars[0], d_cols=T_vars[0], x_cols=selected_X)
ml_l = RandomForestRegressor(**best_params, random_state=42)
ml_m = RandomForestRegressor(**best_params, random_state=42)

# n_rep=10 ensures estimation stability
dml_plr = DoubleMLPLR(dml_data, ml_l, ml_m, n_folds=5, n_rep=10)
dml_plr.fit()

# Pearson correlation of residuals
y_res = dml_plr.residuals['y_res'][:, 0, 0]
d_res = dml_plr.residuals['d_res'][:, 0, 0]
rho, p_val = pearsonr(y_res, d_res)
print(f"Residual Orthogonality: Rho={rho:.4f}, P-val={p_val:.4f} (Closer to 0 is better)")

# ---------- 6. Moderated Mediation Analysis (DML) ----------
print("\n>>> Starting Moderated Mediation Analysis (Causal Forest DML)...")


def run_moderated_mediation(mediator_name):
    print(f"--- Channel: {mediator_name} ---")
    # Step 1: Treatment -> Mediator (Moderated by CEO_Risk_Pref)
    cf_a = CausalForestDML(
        model_y=RandomForestRegressor(**best_params),
        model_t=RandomForestRegressor(**best_params),
        n_estimators=500, n_sequences=10, random_state=42
    )
    cf_a.fit(df[mediator_name].values, df[T_vars[0]].values, X=df[[Moderator]].values, W=df[selected_X].values)

    # Step 2: Mediator -> Outcome (Moderated by CEO_Risk_Pref)
    cf_b = CausalForestDML(
        model_y=RandomForestRegressor(**best_params),
        model_t=RandomForestRegressor(**best_params),
        n_estimators=500, n_sequences=10, random_state=42
    )
    cf_b.fit(df[Y_vars[0]].values, df[mediator_name].values, X=df[[Moderator]].values, W=df[selected_X + T_vars].values)

    return cf_a, cf_b


cf_a_myo, cf_b_myo = run_moderated_mediation('Managerial_Myopia')
cf_a_tra, cf_b_tra = run_moderated_mediation('Transparency')

# ---------- 7. Word Export (APA-Style Table) ----------
doc = Document()
doc.add_heading('Table: Moderated Mediation Results (Robust DML)', level=1)
doc.add_paragraph(
    f"Notes: Fixed Effects (Within-Transformation), Double Selection, and Parameter Tuning applied. Residual Correlation (Rho): {rho:.4f}")

table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, h in enumerate(['Path', 'Mediator', 'ATE', 'T-Value', 'P-Value', 'R-squared']):
    hdr_cells[i].text = h
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True


def add_results_row(tab, path_name, med_name, cf, y_val, x_mod, w_vars):
    ate = cf.ate(df[[x_mod]].values)
    lb, ub = cf.ate_interval(df[[x_mod]].values)
    se = (ub - lb) / (2 * 1.96)
    t_stat = ate / (se + 1e-9)
    p_val = 2 * (1 - norm.cdf(abs(t_stat)))

    # Calculate R-squared for the outcome model
    y_pred = cf.model_y.fit(df[w_vars + [x_mod]], df[y_val]).predict(df[w_vars + [x_mod]])
    r2 = 1 - (((df[y_val] - y_pred) ** 2).sum() / ((df[y_val] - df[y_val].mean()) ** 2).sum())

    row = tab.add_row().cells
    row[0].text, row[1].text = path_name, med_name
    row[2].text, row[3].text, row[4].text, row[5].text = f"{ate:.4f}", f"{t_stat:.2f}", f"{p_val:.4f}", f"{r2:.3f}"


# Populate Table
add_results_row(table, "TR -> Myopia", "Myopia", cf_a_myo, 'Managerial_Myopia', Moderator, selected_X)
add_results_row(table, "Myopia -> GW", "Myopia", cf_b_myo, 'Greenwashing', Moderator, selected_X + T_vars)
add_results_row(table, "TR -> Transp", "Transparency", cf_a_tra, 'Transparency', Moderator, selected_X)
add_results_row(table, "Transp -> GW", "Transparency", cf_b_tra, 'Greenwashing', Moderator, selected_X + T_vars)


# Apply simple formatting to simulate a Three-Line Table
def set_top_bottom_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            # This is a simplified version; APA tables usually have no vertical lines
            pass


output_doc = os.path.join(data_folder, 'DML_Robust_Analysis_EN.docx')
doc.save(output_doc)
print(f"\n>>> Process Complete. Report saved to: {output_doc}")
